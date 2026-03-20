import streamlit as st
from deep_translator import GoogleTranslator
from docx import Document
from docx.oxml.ns import qn
from io import BytesIO
import time
import re

st.set_page_config(
    page_title="WindDoc Translator",
    page_icon="🌬️",
    layout="centered",
    initial_sidebar_state="collapsed",
)

# ======================
# CUSTOM CSS — Dark industrial / wind-energy aesthetic
# Fonts: Bebas Neue (display) + JetBrains Mono (body)
# Palette: near-black bg, electric cyan accent, slate greys
# ======================
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Bebas+Neue&family=JetBrains+Mono:wght@300;400;500;700&display=swap');

/* ── Reset & base ──────────────────────────────────────────── */
html, body, [class*="css"] {
    font-family: 'JetBrains Mono', monospace;
    background-color: #0b0e11;
    color: #c8d6e5;
}
.stApp {
    background: #0b0e11;
}

/* ── Animated background grid ──────────────────────────────── */
.stApp::before {
    content: '';
    position: fixed;
    inset: 0;
    background-image:
        linear-gradient(rgba(0,210,200,0.03) 1px, transparent 1px),
        linear-gradient(90deg, rgba(0,210,200,0.03) 1px, transparent 1px);
    background-size: 48px 48px;
    pointer-events: none;
    z-index: 0;
}

/* ── Hero header ────────────────────────────────────────────── */
.hero {
    text-align: center;
    padding: 3.5rem 1rem 2rem;
    position: relative;
}
.hero-eyebrow {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.82rem;
    font-weight: 500;
    letter-spacing: 0.35em;
    text-transform: uppercase;
    color: #00d2c8;
    margin-bottom: 0.6rem;
}
.hero-title {
    font-family: 'Bebas Neue', sans-serif;
    font-size: clamp(3.2rem, 8vw, 5.5rem);
    letter-spacing: 0.06em;
    line-height: 0.95;
    color: #ffffff;
    margin: 0;
}
.hero-title span {
    color: #00d2c8;
}
.hero-sub {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.82rem;
    color: #4a6080;
    letter-spacing: 0.06em;
    margin-top: 0.9rem;
}
.hero-rule {
    width: 60px;
    height: 2px;
    background: linear-gradient(90deg, #00d2c8, transparent);
    margin: 1.4rem auto 0;
    border: none;
}

/* ── Upload card ────────────────────────────────────────────── */
.card {
    background: #111620;
    border: 1px solid #1e2a3a;
    border-radius: 8px;
    padding: 2rem 2.2rem;
    margin-bottom: 1.2rem;
    position: relative;
    overflow: hidden;
}
.card::before {
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0;
    height: 2px;
    background: linear-gradient(90deg, #00d2c8, #005f8a, transparent);
}
.card-label {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.82rem;
    font-weight: 700;
    letter-spacing: 0.22em;
    text-transform: uppercase;
    color: #00d2c8;
    margin-bottom: 0.8rem;
}

/* ── File uploader ──────────────────────────────────────────── */
[data-testid="stFileUploader"] {
    background: #0d1117 !important;
    border: 1.5px dashed #1e3a50 !important;
    border-radius: 6px !important;
    padding: 1rem !important;
    transition: border-color 0.2s;
}
[data-testid="stFileUploader"]:hover {
    border-color: #00d2c8 !important;
}
[data-testid="stFileUploader"] label {
    color: #4a6080 !important;
    font-size: 0.82rem !important;
    letter-spacing: 0.05em !important;
    text-transform: none !important;
}

/* ── Selectbox ──────────────────────────────────────────────── */
[data-testid="stSelectbox"] > div > div {
    background: #0d1117 !important;
    border: 1px solid #1e2a3a !important;
    border-radius: 4px !important;
    color: #c8d6e5 !important;
    font-family: 'JetBrains Mono', monospace !important;
    font-size: 0.82rem !important;
}
[data-testid="stSelectbox"] > div > div:focus-within {
    border-color: #00d2c8 !important;
    box-shadow: 0 0 0 2px rgba(0,210,200,0.12) !important;
}

/* ── Labels ─────────────────────────────────────────────────── */
label {
    font-family: 'JetBrains Mono', monospace !important;
    font-size: 0.82rem !important;
    letter-spacing: 0.15em !important;
    text-transform: uppercase !important;
    color: #4a6080 !important;
}

/* ── Translate button ───────────────────────────────────────── */
.stButton > button {
    width: 100% !important;
    background: linear-gradient(135deg, #00d2c8 0%, #0099b8 100%) !important;
    color: #0b0e11 !important;
    border: none !important;
    border-radius: 4px !important;
    font-family: 'Bebas Neue', sans-serif !important;
    font-size: 1.15rem !important;
    letter-spacing: 0.18em !important;
    padding: 0.7rem 2rem !important;
    cursor: pointer !important;
    transition: opacity 0.15s, transform 0.1s !important;
    margin-top: 0.4rem !important;
}
.stButton > button:hover {
    opacity: 0.88 !important;
    transform: translateY(-1px) !important;
}
.stButton > button:active {
    transform: translateY(0) !important;
}

/* ── Download button ────────────────────────────────────────── */
.stDownloadButton > button {
    width: 100% !important;
    background: #0b0e11 !important;
    color: #00d2c8 !important;
    border: 1.5px solid #00d2c8 !important;
    border-radius: 4px !important;
    font-family: 'Bebas Neue', sans-serif !important;
    font-size: 1.05rem !important;
    letter-spacing: 0.15em !important;
    padding: 0.6rem 2rem !important;
    transition: background 0.15s !important;
}
.stDownloadButton > button:hover {
    background: rgba(0,210,200,0.08) !important;
}

/* ── Progress bar ───────────────────────────────────────────── */
.stProgress > div > div {
    background: linear-gradient(90deg, #00d2c8, #0099b8) !important;
    border-radius: 2px !important;
}
.stProgress > div {
    background: #1e2a3a !important;
    border-radius: 2px !important;
    height: 4px !important;
}

/* ── Info / success / warning alerts ───────────────────────── */
.stAlert {
    background: #111620 !important;
    border: 1px solid #1e2a3a !important;
    border-radius: 4px !important;
    font-family: 'JetBrains Mono', monospace !important;
    font-size: 0.82rem !important;
    color: #c8d6e5 !important;
}

/* ── Stats row ──────────────────────────────────────────────── */
.stats-row {
    display: flex;
    gap: 1rem;
    margin: 1.4rem 0 0.5rem;
    flex-wrap: wrap;
}
.stat-box {
    flex: 1;
    min-width: 100px;
    background: #0d1117;
    border: 1px solid #1e2a3a;
    border-radius: 6px;
    padding: 1rem 1.2rem;
    text-align: center;
}
.stat-number {
    font-family: 'Bebas Neue', sans-serif;
    font-size: 2rem;
    color: #00d2c8;
    line-height: 1;
}
.stat-label {
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.82rem;
    letter-spacing: 0.18em;
    text-transform: uppercase;
    color: #4a6080;
    margin-top: 0.3rem;
}

/* ── Language badge ─────────────────────────────────────────── */
.lang-badge {
    display: inline-block;
    background: rgba(0,210,200,0.08);
    border: 1px solid rgba(0,210,200,0.25);
    color: #00d2c8;
    font-family: 'JetBrains Mono', monospace;
    font-size: 0.82rem;
    letter-spacing: 0.12em;
    padding: 0.2rem 0.6rem;
    border-radius: 2px;
    margin-bottom: 1rem;
}

/* ── Footer ─────────────────────────────────────────────────── */
.footer {
    text-align: center;
    padding: 2.5rem 0 1rem;
    font-size: 0.82rem;
    letter-spacing: 0.1em;
    color: #1e2a3a;
    text-transform: uppercase;
}

/* ── Hide Streamlit chrome ──────────────────────────────────── */
#MainMenu { visibility: hidden; }
footer { visibility: hidden; }
[data-testid="stToolbar"] { display: none; }
</style>
""", unsafe_allow_html=True)

# ======================
# LANGUAGES
# ======================
languages = {
    "India – Hindi": "hi",
    "India – Tamil": "ta",
    "India – Telugu": "te",
    "India – Kannada": "kn",
    "India – Malayalam": "ml",
    "India – Gujarati": "gu",
    "France – French": "fr",
    "United Kingdom – English": "en",
    "Poland – Polish": "pl",
    "Sweden – Swedish": "sv",
    "Finland – Finnish": "fi",
    "Italy – Italian": "it",
    "Japan – Japanese": "ja",
    "Netherlands – Dutch": "nl",
    "Germany – German": "de",
    "South Korea – Korean": "ko",
    "Australia – English": "en",
    "USA – English": "en",
    "Greece – Greek": "el",
    "Philippines – Filipino": "tl",
    "Egypt – Arabic": "ar",
    "Austria – German": "de",
    "South Africa – Afrikaans": "af",
    "Canada – English": "en",
    "Ireland – Irish (Gaelic)": "ga",
    "Curaçao – Dutch": "nl",
    "Belgium – Dutch": "nl",
    "International Waters – English": "en",
    "Taiwan – Mandarin Chinese": "zh-TW",
    "China – Chinese (Simplified)": "zh-CN",
    "Czech Republic – Czech": "cs",
    "Spain – Spanish": "es",
    "Mexico – Spanish": "es",
    "Brazil – Portuguese": "pt",
    "Turkey – Turkish": "tr",
    "Argentina – Spanish": "es",
    "Lithuania – Lithuanian": "lt",
    "Portugal – Portuguese": "pt",
    "Romania – Romanian": "ro",
    "Cyprus – Greek": "el",
    "Estonia – Estonian": "et",
    "Denmark – Danish": "da",
    "Croatia – Croatian": "hr",
}

# ======================
# WIND INDUSTRY GLOSSARY
# ======================
WIND_GLOSSARY = {
    "es": {
        "cuchillas": "palas", "cuchilla": "pala", "aspas": "palas", "aspa": "pala",
        "paletas": "palas", "paleta": "pala", "hoja": "pala", "hojas": "palas",
        "veletas": "palas", "veleta": "pala", "cabina": "góndola",
        "multiplicador": "multiplicadora", "caja de cambios": "multiplicadora",
        "caja de velocidades": "multiplicadora", "cubo": "buje", "centro": "buje",
        "turbina de viento": "aerogenerador", "molino de viento": "aerogenerador",
        "generador eólico": "aerogenerador", "parque de viento": "parque eólico",
        "granja eólica": "parque eólico", "granja de viento": "parque eólico",
        "control de cabeceo": "control de paso", "control de inclinación": "control de paso",
        "guiñada": "orientación", "puesta en servicio": "puesta en marcha",
        "comisionamiento": "puesta en marcha", "puesta en funcionamiento": "puesta en marcha",
        "terminación mecánica": "finalización mecánica",
        "accidente fatal": "accidente mortal", "cuasi accidente": "cuasi-accidente",
        "casi accidente": "cuasi-accidente", "estación secundaria": "subestación",
        "cable de arreglo": "cable de interconexión", "pilote único": "monopilote",
        "chaqueta": "estructura de celosía", "guindaste": "grúa",
    },
    "pl": {
        "skrzydła": "łopaty", "skrzydło": "łopata", "łopatki": "łopaty", "łopatka": "łopata",
        "wiatrak": "turbina wiatrowa", "park wiatrowy": "farma wiatrowa",
        "farma wiatrakowa": "farma wiatrowa", "skrzynia biegów": "przekładnia",
        "przekładnia zębata": "przekładnia", "kabina": "gondola", "piasta koła": "piasta",
        "oddanie do eksploatacji": "uruchomienie", "kąt skoku": "skok",
        "kąt łopaty": "skok", "ster": "odchylenie",
        "prawie wypadek": "zdarzenie potencjalnie wypadkowe", "dźwig": "żuraw",
    },
    "de": {
        "blätter": "Rotorblätter", "flügel": "Rotorblätter", "rotorflügel": "Rotorblätter",
        "schaufeln": "Rotorblätter", "schaufel": "Rotorblatt",
        "windmühle": "Windkraftanlage", "windrad": "Windkraftanlage",
        "windturbine": "Windkraftanlage", "windfarm": "Windpark",
        "windgenerator": "Windkraftanlage", "kabine": "Gondel", "radnabe": "Nabe",
        "zahnradgetriebe": "Getriebe", "inbetriebsetzung": "Inbetriebnahme",
        "azimutwinkel": "Azimut", "gierwinkel": "Azimut",
        "blattanstellwinkel": "Blattwinkel", "anstellwinkel": "Blattwinkel",
        "beinahunfall": "Beinaheunfall", "einzelpfahl": "Monopfahl",
    },
    "fr": {
        "lames": "pales", "ailes": "pales", "aile": "pale",
        "turbine éolienne": "éolienne", "moulin à vent": "éolienne",
        "ferme éolienne": "parc éolien", "ferme de vent": "parc éolien",
        "boîte de vitesses": "multiplicateur", "transmission": "multiplicateur",
        "cabine": "nacelle", "centre": "moyeu", "commissionning": "mise en service",
        "lacet": "orientation", "pieu unique": "monopieu",
        "accident fatal": "accident mortel", "presque accident": "quasi-accident",
    },
    "it": {
        "lame": "pale", "turbina eolica": "aerogeneratore",
        "mulino a vento": "aerogeneratore", "fattoria eolica": "parco eolico",
        "scatola del cambio": "moltiplicatore", "trasmissione": "moltiplicatore",
        "cabina": "navicella", "messa in funzione": "messa in servizio",
        "incidente fatale": "incidente mortale", "quasi incidente": "quasi-incidente",
    },
    "nl": {
        "bladen": "rotorbladen", "blad": "rotorblad", "vleugels": "rotorbladen",
        "windmolen": "windturbine", "tandwielkast": "versnellingsbak",
        "cabine": "gondel", "spoedregeling": "bladhoekregeling",
        "fataal ongeluk": "dodelijk ongeluk",
    },
    "pt": {
        "lâminas": "pás", "turbina eólica": "aerogerador",
        "moinho de vento": "aerogerador", "fazenda eólica": "parque eólico",
        "caixa de engrenagens": "multiplicadora", "transmissão": "multiplicadora",
        "cabine": "nacele", "posta em serviço": "entrada em operação",
        "comissionamento": "entrada em operação",
    },
    "sv": {
        "blad": "rotorblad", "vingar": "rotorblad", "vinge": "rotorblad",
        "vindturbin": "vindkraftverk", "vindmölla": "vindkraftverk",
        "kugghjulsväxel": "växellåda", "kabin": "gondol",
        "driftsättning": "idrifttagning", "stegkontroll": "bladvinkelreglering",
    },
    "da": {
        "blade": "rotorblade", "vinger": "rotorblade",
        "vindkraftværk": "vindmølle", "vindpark": "vindmøllepark",
        "kabine": "nacelle", "pitchkontrol": "pitchregulering",
    },
    "fi": {
        "lavat": "roottorin lavat", "lapa": "roottorin lapa",
        "tuuliturbiini": "tuulivoimala", "tuulimylly": "tuulivoimala",
        "hajautus": "suuntaus",
    },
    "ja": {
        "刃": "ブレード", "羽根": "ブレード", "翼": "ブレード",
        "風力タービン": "風力発電機", "風車": "風力発電機",
        "風力団地": "ウインドファーム", "ギアボックス": "増速機",
        "変速機": "増速機", "キャビン": "ナセル", "致命的事故": "死亡事故",
        "ニアミス": "ヒヤリハット",
    },
    "ko": {
        "날": "블레이드", "날개": "블레이드", "풍력 터빈": "풍력 발전기",
        "풍차": "풍력 발전기", "풍력 단지": "풍력 발전 단지",
        "기어박스": "증속기", "요": "요잉", "아찔한 순간": "아차 사고",
    },
    "zh-CN": {
        "刀片": "叶片", "刀": "叶片", "桨叶": "叶片",
        "风力涡轮机": "风力发电机", "风车": "风力发电机",
        "风电农场": "风电场", "风能农场": "风电场", "变速箱": "齿轮箱",
        "试运行": "调试", "致命事故": "死亡事故", "险兆": "未遂事故", "吊车": "起重机",
    },
    "zh-TW": {
        "刀片": "葉片", "刀": "葉片", "風力渦輪機": "風力發電機",
        "風車": "風力發電機", "風電農場": "風電場", "試運行": "調試", "致命事故": "死亡事故",
    },
}

PRESERVE_PATTERNS = [
    r'\bIN\.\d{7,12}\b',
    r'\b\d+(?:\.\d+)?\s*(?:MW|kW|m/s|rpm|Hz|kV|MWh|kWh)\b',
    r'\bIEC\s*\d+[-\w]*\b',
    r'\bISO\s*\d+\b',
    r'\bDNV[-\s]\w+\b',
    r'\bVAS\w*\b',
]


def apply_wind_glossary(text, lang):
    for wrong, correct in WIND_GLOSSARY.get(lang, {}).items():
        pattern = re.compile(re.escape(wrong), re.IGNORECASE | re.UNICODE)
        def _replace(m, c=correct):
            return c[0].upper() + c[1:] if m.group(0)[0].isupper() else c
        text = pattern.sub(_replace, text)
    return text


def protect_text(text):
    placeholders = {}
    for pat in PRESERVE_PATTERNS:
        for m in re.finditer(pat, text):
            original = m.group(0)
            if original not in placeholders.values():
                token = f"__PH{len(placeholders)}__"
                placeholders[token] = original
    for token, original in placeholders.items():
        text = text.replace(original, token)
    return text, placeholders


def restore_text(text, placeholders):
    for token, original in placeholders.items():
        text = text.replace(token, original)
    return text


def safe_translate(text, target_lang):
    if not text or text.strip() == "":
        return text
    protected, placeholders = protect_text(text)
    try:
        translated = GoogleTranslator(source="auto", target=target_lang).translate(protected)
        if not translated:
            translated = protected
    except Exception:
        translated = protected
    translated = restore_text(translated, placeholders)
    translated = apply_wind_glossary(translated, target_lang)
    return translated


def run_fmt_key(run):
    try:
        color = run.font.color.rgb if run.font.color and run.font.color.type else None
    except Exception:
        color = None
    return (run.bold, run.italic, run.underline, run.font.size, run.font.name, color)


def translate_paragraph(para, target_lang):
    if not para.runs:
        return
    # Skip paragraphs whose XML contains an embedded image (e.g. logo)
    if element_has_image(para._element):
        return
    groups = []
    for run in para.runs:
        key = run_fmt_key(run)
        if groups and groups[-1][0] == key:
            groups[-1][1].append(run)
        else:
            groups.append((key, [run]))
    for _key, runs in groups:
        combined = "".join(r.text for r in runs)
        if not combined.strip():
            continue
        translated = safe_translate(combined, target_lang)
        runs[0].text = translated
        for r in runs[1:]:
            r.text = ""


def element_has_image(xml_element):
    """Return True if this XML element contains a logo / image node."""
    for tag in ("w:drawing", "v:imagedata", "w:pict", "v:shape"):
        if next(xml_element.iter(qn(tag)), None) is not None:
            return True
    return False


def translate_xml_runs(xml_element, target_lang):
    """Translate all w:t run nodes, skipping any container that holds an image."""
    if element_has_image(xml_element):
        return
    for t_node in xml_element.iter(qn("w:t")):
        original = t_node.text or ""
        if original.strip():
            t_node.text = safe_translate(original, target_lang)


def lock_table_layout(table):
    """
    Freeze every table so translated text cannot shift columns or rows:
      - Sets table layout to 'fixed'  (w:tblLayout type="fixed")
      - Locks each row height to exact (w:trHeight hRule="exact")
      - Locks each cell width to exact (w:tcW type="dxa")
    This prevents longer translated text from pushing the logo cell.
    """
    from lxml import etree

    tbl = table._tbl

    # ── 1. Force fixed table layout ───────────────────────────────────────
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = etree.SubElement(tbl, qn("w:tblPr"))
    tblLayout = tblPr.find(qn("w:tblLayout"))
    if tblLayout is None:
        tblLayout = etree.SubElement(tblPr, qn("w:tblLayout"))
    tblLayout.set(qn("w:type"), "fixed")

    # ── 2. Lock every row height and cell width ───────────────────────────
    for row in table.rows:
        tr = row._tr
        trPr = tr.find(qn("w:trPr"))
        if trPr is None:
            trPr = etree.SubElement(tr, qn("w:trPr"))

        # Lock row height to its current value (exact)
        trHeight = trPr.find(qn("w:trHeight"))
        if trHeight is None:
            # Read current rendered height or fall back to a safe default
            h_val = trHeight.get(qn("w:val"), "567") if trHeight is not None else "567"
            trHeight = etree.SubElement(trPr, qn("w:trHeight"))
            trHeight.set(qn("w:val"), h_val)
        trHeight.set(qn("w:hRule"), "exact")

        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.find(qn("w:tcPr"))
            if tcPr is None:
                tcPr = etree.SubElement(tc, qn("w:tcPr"))

            # Lock cell width to its current value
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is not None:
                tcW.set(qn("w:type"), "dxa")   # twips — fixed unit
            else:
                # Fall back: set a sensible default width
                tcW = etree.SubElement(tcPr, qn("w:tcW"))
                tcW.set(qn("w:w"), "2000")
                tcW.set(qn("w:type"), "dxa")


def count_all_blocks(doc):
    total = len(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                total += len(cell.paragraphs)
    for section in doc.sections:
        for hdr in [section.header, section.footer,
                    section.even_page_header, section.even_page_footer,
                    section.first_page_header, section.first_page_footer]:
            try:
                total += len(hdr.paragraphs)
            except Exception:
                pass
    for txbx in doc.element.iter(qn("w:txbx")):
        total += sum(1 for _ in txbx.iter(qn("w:p")))
    for sdt in doc.element.iter(qn("w:sdt")):
        total += sum(1 for _ in sdt.iter(qn("w:p")))
    return max(total, 1)


def format_eta(seconds):
    return f"{seconds:.1f}s" if seconds < 60 else f"{seconds / 60:.1f}m"


# ======================
# UI — HERO
# ======================
st.markdown("""
<div class="hero">
    <div class="hero-eyebrow">⟡ Wind Energy · HSE Documents</div>
    <div class="hero-title">WIND<span>DOC</span><br>TRANSLATOR</div>
    <div class="hero-sub">Industry-accurate · 40+ languages · Format preserved</div>
    <hr class="hero-rule">
</div>
""", unsafe_allow_html=True)

# ======================
# UI — UPLOAD CARD
# ======================
st.markdown('<div class="card"><div class="card-label">01 — Upload Document</div>', unsafe_allow_html=True)
uploaded_file = st.file_uploader("Drop your DOCX here", type=["docx"], label_visibility="collapsed")
st.markdown('</div>', unsafe_allow_html=True)

# ======================
# UI — LANGUAGE CARD
# ======================
st.markdown('<div class="card"><div class="card-label">02 — Select Target Language</div>', unsafe_allow_html=True)
target_label = st.selectbox("Language", list(languages.keys()), label_visibility="collapsed")
st.markdown('</div>', unsafe_allow_html=True)

# ======================
# UI — TRANSLATE BUTTON
# ======================
run_btn = st.button("▶  TRANSLATE DOCUMENT")

# ======================
# TRANSLATION LOGIC
# ======================
if run_btn and uploaded_file:
    target = languages[target_label]
    doc    = Document(uploaded_file)

    total_blocks = count_all_blocks(doc)
    state        = {"completed": 0}   # mutable dict avoids nonlocal requirement
    start_time   = time.time()

    # Stats row
    st.markdown(f"""
    <div class="stats-row">
        <div class="stat-box">
            <div class="stat-number">{total_blocks}</div>
            <div class="stat-label">Blocks</div>
        </div>
        <div class="stat-box">
            <div class="stat-number">{len(WIND_GLOSSARY.get(target, {}))} </div>
            <div class="stat-label">Glossary Terms</div>
        </div>
        <div class="stat-box">
            <div class="stat-number">{target.upper()}</div>
            <div class="stat-label">Target Lang</div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    progress   = st.progress(0)
    eta_text   = st.empty()
    status_msg = st.empty()
    status_msg.info("⚙ Translating — wind terminology will be corrected automatically…")

    def tick():
        state["completed"] += 1
        c   = state["completed"]
        pct = min(c / total_blocks, 1.0)
        progress.progress(pct)
        elapsed   = time.time() - start_time
        remaining = total_blocks - c
        if c > 0 and remaining > 0:
            eta_text.markdown(
                f'<span style="font-family:JetBrains Mono,monospace;font-size:0.82rem;'
                f'color:#4a6080;">⏳ {int(pct*100)}% · ETA {format_eta((elapsed/c)*remaining)}'
                f' · {c}/{total_blocks} blocks</span>',
                unsafe_allow_html=True,
            )

    # 1. Body paragraphs
    for para in doc.paragraphs:
        translate_paragraph(para, target)
        tick()

    # 2. Body tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    translate_paragraph(para, target)
                    tick()

    # 3. Headers & Footers
    for section in doc.sections:
        for hdr in [section.header, section.footer,
                    section.even_page_header, section.even_page_footer,
                    section.first_page_header, section.first_page_footer]:
            try:
                for para in hdr.paragraphs:
                    translate_paragraph(para, target)
                    tick()
                for table in hdr.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                translate_paragraph(para, target)
                                tick()
            except Exception:
                pass

    # 4. Text boxes (w:txbx) — HSE form fields
    # Each text box is checked individually: boxes containing the logo are skipped,
    # while boxes containing text fields (Date, From, Re: etc.) are translated.
    for txbx in doc.element.iter(qn("w:txbx")):
        if not element_has_image(txbx):
            translate_xml_runs(txbx, target)
        for _ in txbx.iter(qn("w:p")):
            tick()

    # 5. Content controls (w:sdt) — structured fields
    for sdt in doc.element.iter(qn("w:sdt")):
        if not element_has_image(sdt):
            translate_xml_runs(sdt, target)
        for _ in sdt.iter(qn("w:p")):
            tick()

    # 6. Lock all table layouts so translated text cannot shift the logo cell
    for table in doc.tables:
        lock_table_layout(table)
    for section in doc.sections:
        for hdr in [section.header, section.footer,
                    section.even_page_header, section.even_page_footer,
                    section.first_page_header, section.first_page_footer]:
            try:
                for table in hdr.tables:
                    lock_table_layout(table)
            except Exception:
                pass

    # Save
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    progress.progress(1.0)
    eta_text.empty()
    status_msg.success("✓ Translation complete — wind terminology corrected")

    safe_name = re.sub(r'[^\w\-]', '_', target_label)

    st.markdown('<div style="margin-top:1rem;">', unsafe_allow_html=True)
    st.download_button(
        "⬇  DOWNLOAD TRANSLATED DOCX",
        data=output,
        file_name=f"translated_{safe_name}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    st.markdown('</div>', unsafe_allow_html=True)

elif run_btn and not uploaded_file:
    st.warning("⚠ Please upload a DOCX file first.")

# ======================
# FOOTER
# ======================
st.markdown("""
<div class="footer">
    WindDoc Translator · Wind Energy HSE · Powered by Google Translate + Industry Glossary
</div>
""", unsafe_allow_html=True)
